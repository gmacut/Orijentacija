"""Generira kontrolne kartice tečajaca za vježbu orijentacije.

Izlaz: Kontrolne_kartice_za_tečajce_Kalnik.docx
       Kontrolne_kartice_za_tečajce_Okić.docx
Pokretanje: python tools/generate_kartice.py  (iz korijena repoa)
"""

import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Kalnik ───────────────────────────────────────────────────────────────────

KALNIK_KT_OPIS = {
    "DOM":     "Planinarski dom",
    "KT LIA":  "Raskrižje ceste i planinske staze",
    "KT LJIA": "Vrh",
    "KT LJA":  "Raskrižje ceste i planinske staze",
    "KT LZA":  "Raskrižje ceste i planinske staze",
    "KT LSIA": "Raskrižje ceste i planinarskih staza",
    "KT DJZA": "Raskrižje planinarskih staza",
    "KT DIA":  "Raskrižje planinarskih staza",
    "KT DSZA": "Raskrižje planinarskih staza",
    "KT DZA":  "Raskrižje planinarskih staza",
}

KALNIK_SEGMENTI = {
    ("DOM",     "KT LIA"):  (261,  252),
    ("DOM",     "KT LSIA"): (302,  488),
    ("DOM",     "KT DJZA"): ( 91,  314),
    ("DOM",     "KT DZA"):  ( 60,  282),
    ("KT LIA",  "DOM"):     ( 81,  252),
    ("KT LIA",  "KT LJIA"): (258,  399),
    ("KT LIA",  "KT DJZA"): ( 86,  564),
    ("KT LIA",  "KT DZA"):  ( 70,  525),
    ("KT LJIA", "KT LIA"):  ( 78,  399),
    ("KT LJIA", "KT LZA"):  (270,  860),
    ("KT LZA",  "KT LJIA"): ( 90,  860),
    ("KT LZA",  "KT LSIA"): ( 71, 1150),
    ("KT LSIA", "DOM"):     (122,  488),
    ("KT LSIA", "KT LZA"):  (251, 1150),
    ("KT LSIA", "KT DJZA"): (110,  775),
    ("KT LSIA", "KT DZA"):  (100,  669),
    ("KT DJZA", "DOM"):     (271,  314),
    ("KT DJZA", "KT LIA"):  (266,  564),
    ("KT DJZA", "KT LSIA"): (290,  775),
    ("KT DJZA", "KT DIA"):  ( 71,  465),
    ("KT DIA",  "KT DJZA"): (251,  465),
    ("KT DIA",  "KT DSZA"): (298,  360),
    ("KT DSZA", "KT DIA"):  (118,  360),
    ("KT DSZA", "KT DZA"):  (228,  259),
    ("KT DZA",  "DOM"):     (240,  282),
    ("KT DZA",  "KT LIA"):  (250,  525),
    ("KT DZA",  "KT LSIA"): (280,  669),
    ("KT DZA",  "KT DSZA"): ( 48,  259),
}

KALNIK_GRUPE = {
    "G1": ["DOM", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "DOM"],
    "G2": ["DOM", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "DOM"],
    "G3": ["DOM", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "DOM"],
    "G4": ["DOM", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "DOM"],
    "G5": ["DOM", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "DOM"],
    "G6": ["DOM", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "DOM"],
    "G7": ["DOM", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "DOM"],
    "G8": ["DOM", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "DOM"],
}

# ── Okić ─────────────────────────────────────────────────────────────────────

OKIC_KT_OPIS = {
    "DOM":    "Planinarski dom",
    "KT JA":  "Raskrižje ceste i planinske staze",
    "KT JZA": "Raskrižje planinske staze i potoka",
    "KT ZA":  "Raskrižje puta i ceste",
    "KT SZA": "Raskrižje planinske staze i ceste",
    "KT SA":  "Raskrižje planinske staze i puta",
    "KT SIA": "Raskrižje ceste i puta",
    "KT SIB": "Raskrižje dviju cesta",
    "KT JIA": "Prelazak ceste preko potoka",
}

OKIC_SEGMENTI = {
    ("DOM",    "KT JA"):  (308, 236),
    ("DOM",    "KT JZA"): (310, 611),
    ("DOM",    "KT ZA"):  (301, 989),
    ("DOM",    "KT SIB"): (  6, 981),
    ("DOM",    "KT JIA"): (358, 500),
    ("KT JA",  "KT JZA"): (312, 375),
    ("KT JA",  "KT ZA"):  (299, 755),
    ("KT JA",  "DOM"):    (128, 236),
    ("KT JZA", "KT SZA"): (  7, 412),
    ("KT JZA", "KT SA"):  (  1, 754),
    ("KT JZA", "KT JA"):  (132, 375),
    ("KT JZA", "DOM"):    (130, 611),
    ("KT ZA",  "KT SZA"): ( 56, 519),
    ("KT ZA",  "KT SIA"): ( 58, 874),
    ("KT ZA",  "KT JA"):  (119, 755),
    ("KT ZA",  "DOM"):    (121, 989),
    ("KT SZA", "KT SA"):  (353, 348),
    ("KT SZA", "KT SIA"): ( 60, 355),
    ("KT SZA", "KT JZA"): (187, 412),
    ("KT SZA", "KT ZA"):  (236, 519),
    ("KT SA",  "KT SIB"): (107, 587),
    ("KT SA",  "KT JZA"): (181, 754),
    ("KT SA",  "KT SZA"): (173, 348),
    ("KT SIA", "KT JIA"): (170, 487),
    ("KT SIA", "KT SZA"): (240, 355),
    ("KT SIA", "KT ZA"):  (238, 874),
    ("KT SIB", "DOM"):    (186, 981),
    ("KT SIB", "KT SA"):  (287, 587),
    ("KT JIA", "DOM"):    (178, 500),
    ("KT JIA", "KT SIA"): (350, 487),
}

OKIC_GRUPE = {
    "G1": ["DOM", "KT JA",  "KT JZA", "KT SA",  "KT SIB", "DOM"],
    "G2": ["DOM", "KT JZA", "KT SZA", "KT SA",  "KT SIB", "DOM"],
    "G3": ["DOM", "KT ZA",  "KT SZA", "KT SIA", "KT JIA", "DOM"],
    "G4": ["DOM", "KT JA",  "KT ZA",  "KT SIA", "KT JIA", "DOM"],
    "G5": ["DOM", "KT SIB", "KT SA",  "KT JZA", "KT JA",  "DOM"],
    "G6": ["DOM", "KT SIB", "KT SA",  "KT SZA", "KT JZA", "DOM"],
    "G7": ["DOM", "KT JIA", "KT SIA", "KT SZA", "KT ZA",  "DOM"],
    "G8": ["DOM", "KT JIA", "KT SIA", "KT ZA",  "KT JA",  "DOM"],
}

# ── Upute (zajedničke za sve lokacije) ───────────────────────────────────────

UPUTE = [
    "Staza kreće i završava na DOM-u (planinarskom domu).",
    "Koristeći tablicu segmenata i kompas, ucrtaj sve kontrolne točke na karti. "
    "Voditelj vježbe mora potvrditi položaj svih točaka prije polaska na teren.",
    "Na svakom KT-u dokaži prolazak bušenjem kartona bušačem koji se nalazi na toj točki.",
    "Svaki puta kada se buši karton, druga osoba iz grupe SMS-om dojavuje voditelju vježbe "
    "trenutne GPS koordinate. Koordinate su dostupne u navigacijskoj aplikaciji na mobitelu.",
    "Na terenu je dopušteno koristiti elektroničku navigaciju (mobitel, GPS aplikacija).",
    "Kretati se smije samo po označenim planinarskim stazama, cestama i šumskim putovima.",
]

# ── Pomoćne funkcije ─────────────────────────────────────────────────────────

def cell_borders(cell, sz="12", color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def write_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_para(doc, text="", bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    p.runs[0]._r.append(br) if p.runs else p.add_run()._r.append(br)


# ── Generiranje dokumenta ────────────────────────────────────────────────────

def generate_doc(output_path, title, grupe, segmenti, kt_opis, landscape=True):
    doc = Document()

    # Postavke stranice
    section = doc.sections[0]
    if landscape:
        section.page_width, section.page_height = Cm(29.7), Cm(21.0)
    else:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    usable_w = section.page_width - section.left_margin - section.right_margin

    # Širine stupaca tablice segmenata — opis dobiva ostatak
    if landscape:
        seg_fixed = [Cm(0.8), Cm(3.2), Cm(1.8), Cm(2.4), Cm(3.2)]
    else:
        seg_fixed = [Cm(0.7), Cm(2.8), Cm(1.6), Cm(2.2), Cm(2.8)]
    seg_widths = seg_fixed + [usable_w - sum(seg_fixed)]

    # Ukloni zadani prazan paragraf koji Word automatski stvara
    body = doc.element.body
    default_p = body.find(qn("w:p"))
    if default_p is not None:
        body.remove(default_p)

    for i, (grupa, ruta) in enumerate(grupe.items()):
        if i > 0:
            add_page_break(doc)

        # ── Bušilica tablica — na samom vrhu stranice ──────────────────────
        punch_pts = ruta[:-1]   # DOM + KT-ovi; bez povratnog DOM-a na kraju
        n_cols = len(punch_pts)
        col_w = usable_w // n_cols

        tbl = doc.add_table(rows=2, cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        tbl.rows[0].height = Cm(3.2)
        tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        tbl.rows[1].height = Cm(0.7)
        tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for j, kt in enumerate(punch_pts):
            c0 = tbl.cell(0, j)
            c0.width = col_w
            cell_borders(c0, sz="18")

            c1 = tbl.cell(1, j)
            c1.width = col_w
            cell_borders(c1, sz="12")
            cell_shading(c1, "DDDDDD")
            write_cell(c1, kt, bold=True, size=8)

        # ── Naslov i grupa (ispod tablice) ─────────────────────────────────
        add_para(doc, title, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=8, space_after=0)
        add_para(doc, f"Grupa {i + 1}", bold=True, size=15,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=6)

        # ── Upute ──────────────────────────────────────────────────────────
        add_para(doc, "Upute:", bold=True, size=9, space_before=0, space_after=2)
        for u in UPUTE:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(u)
            run.font.size = Pt(9)

        # ── Tablica segmenata ───────────────────────────────────────────────
        add_para(doc, "Segmenti rute:", bold=True, size=9,
                 space_before=8, space_after=2)

        n_seg = len(ruta) - 1
        stbl = doc.add_table(rows=n_seg + 1, cols=6)
        stbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        stbl.style = "Table Grid"

        hdrs = ["#", "Od", "Azimut", "Udaljenost", "Do", "Opis odredišne KT"]
        for k, (h_txt, w) in enumerate(zip(hdrs, seg_widths)):
            c = stbl.cell(0, k)
            c.width = w
            cell_shading(c, "BBBBBB")
            write_cell(c, h_txt, bold=True, size=9)

        row_aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ]
        for s in range(n_seg):
            od, do = ruta[s], ruta[s + 1]
            az, ud = segmenti[(od, do)]
            row_data = [str(s + 1), od, f"{az}°", f"{ud} m", do, kt_opis[do]]
            for k, (txt, w, align) in enumerate(zip(row_data, seg_widths, row_aligns)):
                c = stbl.cell(s + 1, k)
                c.width = w
                write_cell(c, txt, size=9, align=align)

    doc.save(output_path)
    print(f"Spremljeno: {output_path}")


# ── Glavni program ────────────────────────────────────────────────────────────

def main():
    generate_doc(
        "Kalnik/Kontrolne_kartice_za_tecajce_Kalnik.docx",
        "Vježba orijentacije — Kalnik",
        KALNIK_GRUPE, KALNIK_SEGMENTI, KALNIK_KT_OPIS,
        landscape=True,
    )
    generate_doc(
        "Okic/Kontrolne_kartice_za_tecajce_Okic.docx",
        "Vježba orijentacije — Okić",
        OKIC_GRUPE, OKIC_SEGMENTI, OKIC_KT_OPIS,
        landscape=False,
    )


if __name__ == "__main__":
    main()
