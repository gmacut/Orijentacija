"""Generira materijale za organizatora vježbe orijentacije.

Izlaz: Materijali_za_organizatora_Kalnik.docx
       Materijali_za_organizatora_Okić.docx
Pokretanje: python tools/generate_organizator.py  (iz korijena repoa)
"""

import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Kalnik — KT podaci (naziv, puni opis, DMS, decimalno) ───────────────────

KALNIK_KT_DATA = {
    "DOM":     ("Terasa planinarskog doma Kalnik",
                '46°7\'56.55"N 16°27\'46.67"E', "46.132375, 16.462964"),
    "KT LIA":  ("Raskrižje ceste i markirane pl. staze bez broja",
                '46°7\'55.23"N 16°27\'35.05"E', "46.132008, 16.459736"),
    "KT LJIA": ("Vrh Vranilac",
                '46°7\'52.60"N 16°27\'16.82"E', "46.131278, 16.454672"),
    "KT LJA":  ("Raskrižje ceste i pl. staze 115 (alt. za KT LJIA)",
                '46°7\'48.80"N 16°27\'29.36"E', "46.130222, 16.458156"),
    "KT LZA":  ("Raskrižje ceste i pl. staze 101",
                '46°7\'52.47"N 16°26\'36.64"E', "46.131242, 16.443511"),
    "KT LSIA": ("Raskrižje ceste i pl. staza 102 i 103",
                '46°8\'4.89"N 16°27\'27.29"E',  "46.134692, 16.457581"),
    "KT DJZA": ("Raskrižje pl. staza 152 i 126",
                '46°7\'56.38"N 16°28\'1.36"E',  "46.132328, 16.467044"),
    "KT DIA":  ("Raskrižje pl. staza 152, 147 i 150",
                '46°8\'1.19"N 16°28\'21.93"E',  "46.133664, 16.472758"),
    "KT DSZA": ("Raskrižje pl. staza 123 i markirane pl. staze bez broja",
                '46°8\'6.69"N 16°28\'7.08"E',   "46.135192, 16.468633"),
    "KT DZA":  ("Raskrižje pl. staza 123 i 127",
                '46°8\'1.11"N 16°27\'58.07"E',  "46.133642, 16.466131"),
}

KALNIK_GRUPE = {
    "G1": ["DOM", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "DOM"],
    "G2": ["DOM", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "DOM"],
    "G3": ["DOM", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "DOM"],
    "G4": ["DOM", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "DOM"],
    "G5": ["DOM", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "DOM"],
    "G6": ["DOM", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "DOM"],
    "G7": ["DOM", "KT LSIA", "KT LZA", "KT LJIA", "KT LIA", "KT DZA", "KT DSZA", "KT DIA", "KT DJZA", "DOM"],
    "G8": ["DOM", "KT DJZA", "KT DIA", "KT DSZA", "KT DZA", "KT LIA", "KT LJIA", "KT LZA", "KT LSIA", "DOM"],
}

# ── Okić — KT podaci ─────────────────────────────────────────────────────────

OKIC_KT_DATA = {
    "DOM":    ("Cesta ispred planinarskog doma",
               '45°44\'56.13"N 15°42\'11.66"E', "45.748925, 15.703239"),
    "KT JA":  ("Raskrižje ceste i plan. staze 1",
               '45°45\'0.86"N 15°42\'3.07"E',   "45.750239, 15.700853"),
    "KT JZA": ("Raskrižje plan. staze 1 i potoka",
               '45°45\'8.91"N 15°41\'50.05"E',  "45.752475, 15.697236"),
    "KT ZA":  ("Raskrižje puta i ceste",
               '45°45\'12.79"N 15°41\'32.49"E', "45.753553, 15.692358"),
    "KT SZA": ("Raskrižje plan. staze 1 i ceste",
               '45°45\'22.13"N 15°41\'52.52"E', "45.756147, 15.697922"),
    "KT SA":  ("Raskrižje plan. staze 1 i puta",
               '45°45\'33.31"N 15°41\'50.66"E', "45.759253, 15.697406"),
    "KT SIA": ("Raskrižje ceste i puta",
               '45°45\'27.84"N 15°42\'6.83"E',  "45.757733, 15.701897"),
    "KT SIB": ("Raskrižje dvije ceste",
               '45°45\'27.70"N 15°42\'16.67"E', "45.757694, 15.704631"),
    "KT JIA": ("Prelazak ceste preko potoka",
               '45°45\'12.32"N 15°42\'10.90"E', "45.753422, 15.703028"),
}

OKIC_GRUPE = {
    "G1": ["DOM", "KT JA", "KT JZA", "KT SA", "KT SIB", "DOM"],
    "G2": ["DOM", "KT JZA", "KT SZA", "KT SA", "KT SIB", "DOM"],
    "G3": ["DOM", "KT ZA", "KT SZA", "KT SIA", "KT JIA", "DOM"],
    "G4": ["DOM", "KT JA", "KT ZA", "KT SIA", "KT JIA", "DOM"],
}

# ── Pomoćne funkcije ─────────────────────────────────────────────────────────

def cell_borders_full(cell, sz="12", color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def cell_border_bottom_only(cell, sz="6", color="888888"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tcPr.append(borders)


def cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def write_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_para(doc, text="", bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── Generiranje dokumenta ────────────────────────────────────────────────────

def generate_doc(output_path, title, grupe, kt_data):
    doc = Document()

    # Portrait A4, uske margine
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    usable_w = section.page_width - section.left_margin - section.right_margin  # 18 cm

    # Širine stupaca tablice KT-ova
    col_widths = [Cm(2.0), Cm(5.5), Cm(4.3), Cm(3.2), Cm(3.0)]  # = 18 cm

    # Ukloni zadani prazan paragraf
    body = doc.element.body
    default_p = body.find(qn("w:p"))
    if default_p is not None:
        body.remove(default_p)

    for i, (grupa, ruta) in enumerate(grupe.items()):
        if i > 0:
            add_page_break(doc)

        # ── Zaglavlje ──────────────────────────────────────────────────────
        add_para(doc, title, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=0, space_after=2)
        add_para(doc, f"Grupa {i + 1}", bold=True, size=20,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)

        route_str = " → ".join(ruta)
        add_para(doc, f"Ruta: {route_str}", size=9,
                 space_before=0, space_after=10)

        # ── Prostor za polaznike ───────────────────────────────────────────
        add_para(doc, "Polaznici:", bold=True, size=10,
                 space_before=0, space_after=4)

        # Tablica s pet redaka, samo donji obrub (linije za pisanje)
        member_tbl = doc.add_table(rows=5, cols=1)
        member_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for row in member_tbl.rows:
            row.height = Cm(0.75)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            cell = row.cells[0]
            cell.width = usable_w
            cell_border_bottom_only(cell)

        # ── Tablica kontrolnih točaka ──────────────────────────────────────
        add_para(doc, "Kontrolne točke:", bold=True, size=10,
                 space_before=10, space_after=3)

        visit_kts = ruta[:-1]   # DOM + KT-ovi; bez povratnog DOM-a
        n_rows = len(visit_kts) + 1  # + zaglavlje

        kt_tbl = doc.add_table(rows=n_rows, cols=5)
        kt_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        kt_tbl.style = "Table Grid"

        hdrs = ["KT", "Opis", "Koordinate (DMS)", "Koordinate (dec.)", "Dolazak (hh:mm)"]
        h_aligns = [WD_ALIGN_PARAGRAPH.CENTER] * 5
        for k, (h_txt, w) in enumerate(zip(hdrs, col_widths)):
            c = kt_tbl.cell(0, k)
            c.width = w
            cell_shading(c, "BBBBBB")
            write_cell(c, h_txt, bold=True, size=9, align=h_aligns[k])

        row_aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ]
        for row_idx, kt in enumerate(visit_kts):
            opis, dms, dec = kt_data[kt]
            row_data = [kt, opis, dms, dec, ""]
            for k, (txt, w, align) in enumerate(zip(row_data, col_widths, row_aligns)):
                c = kt_tbl.cell(row_idx + 1, k)
                c.width = w
                write_cell(c, txt, size=9, align=align)

    # ── List za referencu bušenja ──────────────────────────────────────────────
    add_page_break(doc)
    add_para(doc, "Referenca bušenja — postavljanje poligona", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)
    add_para(doc, title, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=6)
    add_para(doc, "Kod postavljanja poligona probušite ovaj papir bušačem svake kontrolne točke "
             "kako biste imali referencu uzoraka bušenja po KT-u.",
             size=9, space_before=0, space_after=8)

    # Sve KT-ove osim DOM-a (DOM nema bušač)
    punch_kts = [(kt, data) for kt, data in kt_data.items() if kt != "DOM"]
    ref_col_widths = [Cm(2.5), Cm(2.0), Cm(5.5), Cm(4.0), Cm(4.0)]  # = 18 cm

    ref_tbl = doc.add_table(rows=len(punch_kts) + 1, cols=5)
    ref_tbl.style = "Table Grid"

    ref_hdrs = ["Uzorak bušenja", "KT", "Opis", "Koordinate (DMS)", "Koordinate (dec.)"]
    for k, (h, w) in enumerate(zip(ref_hdrs, ref_col_widths)):
        c = ref_tbl.cell(0, k)
        c.width = w
        cell_shading(c, "BBBBBB")
        write_cell(c, h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    ref_row_aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]
    for row_idx, (kt, (opis, dms, dec)) in enumerate(punch_kts):
        row = ref_tbl.rows[row_idx + 1]
        row.height = Cm(2.2)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for k, (txt, w, align) in enumerate(
            zip(["", kt, opis, dms, dec], ref_col_widths, ref_row_aligns)
        ):
            c = ref_tbl.cell(row_idx + 1, k)
            c.width = w
            write_cell(c, txt, size=9, align=align)

    doc.save(output_path)
    print(f"Spremljeno: {output_path}")


# ── Glavni program ────────────────────────────────────────────────────────────

def main():
    generate_doc(
        "Kalnik/Materijali_za_organizatora_Kalnik.docx",
        "Vježba orijentacije — Kalnik",
        KALNIK_GRUPE, KALNIK_KT_DATA,
    )
    generate_doc(
        "Okic/Materijali_za_organizatora_Okic.docx",
        "Vježba orijentacije — Okić",
        OKIC_GRUPE, OKIC_KT_DATA,
    )


if __name__ == "__main__":
    main()
